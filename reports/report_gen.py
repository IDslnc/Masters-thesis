from docx import Document
from docx.shared import Inches

def generate_report(patient_name, image_path, result_text, save_path):
    doc = Document()
    doc.add_heading('Отчёт по анализу снимка', 0)
    doc.add_paragraph(f'Пациент: {patient_name}')
    doc.add_paragraph(f'Результат анализа: {result_text}')
    doc.add_picture(image_path, width=Inches(5.5))
    doc.save(save_path)
